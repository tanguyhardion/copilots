import json
import sys
from datetime import datetime
from dateutil.relativedelta import relativedelta
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.opc.constants import RELATIONSHIP_TYPE as RT

# ─────────────────────────────────────────────
# CONSTANTS
# ─────────────────────────────────────────────

FONT_NAME = "Arial"
COLOR_BLACK = RGBColor(0x00, 0x00, 0x00)
COLOR_BLUE_DARK = RGBColor(0x0E, 0x41, 0x94)
COLOR_BLUE_LIGHT = RGBColor(0x15, 0x93, 0xCB)
COLOR_BLUE_MID = RGBColor(0x39, 0x8E, 0xC9)
COLOR_HEADER_FILL = RGBColor(0x15, 0x60, 0x82)
COLOR_GRAY = RGBColor(0xBF, 0xBF, 0xBF)

COL_LEFT_CM = 5.0
COL_RIGHT_CM = 13.3
TABLE_INDENT_CM = COL_LEFT_CM  # Single source of truth for all indented tables/dividers
FULL_WIDTH_CM = 18.3
DEFAULT_SIZE = 9
SECTION_SPACE = 8
SKILL_SUBSECTION_SPACE = 12  # Extra spacing between personal skills subsections
BULLET_INDENT_CM = 0.4
BULLET_CHAR = "\u25cf"


# ─────────────────────────────────────────────
# LOW-LEVEL XML HELPERS
# ─────────────────────────────────────────────


def set_cell_margins(cell, top=0, bottom=0, left=0.19, right=0.19):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    for existing in tcPr.findall(qn("w:tcMar")):
        tcPr.remove(existing)
    mar = OxmlElement("w:tcMar")
    for side, val in [
        ("top", top),
        ("bottom", bottom),
        ("left", left),
        ("right", right),
    ]:
        el = OxmlElement(f"w:{side}")
        el.set(qn("w:w"), str(int(val * 567)))
        el.set(qn("w:type"), "dxa")
        mar.append(el)
    tcPr.append(mar)


def set_cell_shading(cell, fill_hex: str):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    for existing in tcPr.findall(qn("w:shd")):
        tcPr.remove(existing)
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill_hex.lstrip("#"))
    tcPr.append(shd)


def set_cell_border(cell, **kwargs):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    for existing in tcPr.findall(qn("w:tcBorders")):
        tcPr.remove(existing)
    tcBorders = OxmlElement("w:tcBorders")
    for side, cfg in kwargs.items():
        el = OxmlElement(f"w:{side}")
        el.set(qn("w:val"), cfg.get("val", "single"))
        el.set(qn("w:sz"), str(cfg.get("sz", 8)))
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), cfg.get("color", "000000"))
        tcBorders.append(el)
    tcPr.append(tcBorders)


def clear_cell_borders(cell):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    for existing in tcPr.findall(qn("w:tcBorders")):
        tcPr.remove(existing)
    tcBorders = OxmlElement("w:tcBorders")
    for side in ["top", "left", "bottom", "right", "insideH", "insideV"]:
        el = OxmlElement(f"w:{side}")
        el.set(qn("w:val"), "none")
        tcBorders.append(el)
    tcPr.append(tcBorders)


def set_table_borders_none(table):
    tbl = table._tbl
    tblPr = tbl.find(qn("w:tblPr"))
    if tblPr is None:
        tblPr = OxmlElement("w:tblPr")
        tbl.insert(0, tblPr)
    for existing in tblPr.findall(qn("w:tblBorders")):
        tblPr.remove(existing)
    tblBorders = OxmlElement("w:tblBorders")
    for side in ["top", "left", "bottom", "right", "insideH", "insideV"]:
        el = OxmlElement(f"w:{side}")
        el.set(qn("w:val"), "none")
        tblBorders.append(el)
    tblPr.append(tblBorders)


def apply_standard_table_borders(table):
    for row in table.rows:
        for cell in row.cells:
            set_cell_border(
                cell,
                top={"sz": 8, "val": "single", "color": "0E4194"},
                bottom={"sz": 8, "val": "single", "color": "0E4194"},
                left={"sz": 8, "val": "single", "color": "0E4194"},
                right={"sz": 8, "val": "single", "color": "0E4194"},
            )


def set_row_height(row, height_cm: float, exact: bool = True):
    tr = row._tr
    trPr = tr.get_or_add_trPr()
    for existing in trPr.findall(qn("w:trHeight")):
        trPr.remove(existing)
    trH = OxmlElement("w:trHeight")
    trH.set(qn("w:val"), str(int(height_cm * 567)))
    trH.set(qn("w:hRule"), "exact" if exact else "atLeast")
    trPr.append(trH)


def set_cell_valign(cell, val="bottom"):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    for existing in tcPr.findall(qn("w:vAlign")):
        tcPr.remove(existing)
    vAlign = OxmlElement("w:vAlign")
    vAlign.set(qn("w:val"), val)
    tcPr.append(vAlign)


def set_para_spacing(para, before_pt=0, after_pt=0, line_pt=None):
    pPr = para._p.get_or_add_pPr()
    for existing in pPr.findall(qn("w:spacing")):
        pPr.remove(existing)
    spc = OxmlElement("w:spacing")
    spc.set(qn("w:before"), str(int(before_pt * 20)))
    spc.set(qn("w:after"), str(int(after_pt * 20)))
    if line_pt is not None:
        spc.set(qn("w:line"), str(int(line_pt * 20)))
        spc.set(qn("w:lineRule"), "exact")
    pPr.append(spc)


def add_tab_stop(para, position_cm: float, align: str = "left"):
    pPr = para._p.get_or_add_pPr()
    tabs = pPr.find(qn("w:tabs"))
    if tabs is None:
        tabs = OxmlElement("w:tabs")
        pPr.append(tabs)
    tab = OxmlElement("w:tab")
    tab.set(qn("w:val"), align)
    tab.set(qn("w:pos"), str(int(position_cm * 567)))
    tabs.append(tab)


def set_para_indent(para, left_cm=0.0, hanging_cm=0.0):
    pPr = para._p.get_or_add_pPr()
    for existing in pPr.findall(qn("w:ind")):
        pPr.remove(existing)
    ind = OxmlElement("w:ind")
    ind.set(qn("w:left"), str(int(left_cm * 567)))
    ind.set(qn("w:hanging"), str(int(hanging_cm * 567)))
    pPr.append(ind)


def set_col_width(cell, width_cm: float):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    for existing in tcPr.findall(qn("w:tcW")):
        tcPr.remove(existing)
    tcW = OxmlElement("w:tcW")
    tcW.set(qn("w:w"), str(int(width_cm * 567)))
    tcW.set(qn("w:type"), "dxa")
    tcPr.append(tcW)


def indent_table(table, indent_cm: float):
    tblPr = table._tbl.find(qn("w:tblPr"))
    if tblPr is None:
        tblPr = OxmlElement("w:tblPr")
        table._tbl.insert(0, tblPr)
    for existing in tblPr.findall(qn("w:tblInd")):
        tblPr.remove(existing)
    for existing in tblPr.findall(qn("w:tblW")):
        tblPr.remove(existing)
    tblInd = OxmlElement("w:tblInd")
    tblInd.set(qn("w:w"), str(int(indent_cm * 567)))
    tblInd.set(qn("w:type"), "dxa")
    tblPr.append(tblInd)
    tblW = OxmlElement("w:tblW")
    tblW.set(qn("w:w"), "0")
    tblW.set(qn("w:type"), "auto")
    tblPr.append(tblW)


def set_table_width(table, width_cm: float):
    tblPr = table._tbl.find(qn("w:tblPr"))
    if tblPr is None:
        tblPr = OxmlElement("w:tblPr")
        table._tbl.insert(0, tblPr)
    for existing in tblPr.findall(qn("w:tblW")):
        tblPr.remove(existing)
    tblW = OxmlElement("w:tblW")
    tblW.set(qn("w:w"), str(int(width_cm * 567)))
    tblW.set(qn("w:type"), "dxa")
    tblPr.append(tblW)


def add_hyperlink(para, text: str, url: str, size=DEFAULT_SIZE, color=COLOR_BLUE_LIGHT):
    part = para.part
    r_id = part.relate_to(url, RT.HYPERLINK, is_external=True)
    hl = OxmlElement("w:hyperlink")
    hl.set(qn("r:id"), r_id)
    r = OxmlElement("w:r")
    rPr = OxmlElement("w:rPr")
    rFonts = OxmlElement("w:rFonts")
    rFonts.set(qn("w:ascii"), FONT_NAME)
    rFonts.set(qn("w:hAnsi"), FONT_NAME)
    sz = OxmlElement("w:sz")
    sz.set(qn("w:val"), str(int(size * 2)))
    clr = OxmlElement("w:color")
    hex_color = "%02X%02X%02X" % (color[0], color[1], color[2])
    clr.set(qn("w:val"), hex_color)
    u = OxmlElement("w:u")
    u.set(qn("w:val"), "single")
    rPr.append(rFonts)
    rPr.append(sz)
    rPr.append(clr)
    rPr.append(u)
    r.append(rPr)
    t = OxmlElement("w:t")
    t.text = text
    r.append(t)
    hl.append(r)
    para._p.append(hl)


# ─────────────────────────────────────────────
# RUN / PARAGRAPH HELPERS
# ─────────────────────────────────────────────


def apply_run_fmt(
    run, size=DEFAULT_SIZE, bold=False, color=COLOR_BLACK, underline=False
):
    run.font.name = FONT_NAME
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.color.rgb = color
    run.font.underline = underline


def add_run(
    para, text, size=DEFAULT_SIZE, bold=False, color=COLOR_BLACK, underline=False
):
    run = para.add_run(text)
    apply_run_fmt(run, size=size, bold=bold, color=color, underline=underline)
    return run


def new_para(doc, alignment=WD_ALIGN_PARAGRAPH.LEFT, before_pt=0, after_pt=0):
    para = doc.add_paragraph()
    para.alignment = alignment
    set_para_spacing(para, before_pt, after_pt)
    return para


# ─────────────────────────────────────────────
# DATE HELPERS
# ─────────────────────────────────────────────

MONTHS = [
    "Jan",
    "Feb",
    "Mar",
    "Apr",
    "May",
    "Jun",
    "Jul",
    "Aug",
    "Sep",
    "Oct",
    "Nov",
    "Dec",
]


def fmt_date(ym: str) -> str:
    if ym == "Present":
        return "Present"
    y, m = ym.split("-")
    return f"{MONTHS[int(m) - 1]} {y}"


def fmt_date_range(date_from: str, date_to: str) -> str:
    return f"{fmt_date(date_from)} – {fmt_date(date_to)}"


def calc_months(date_from: str, date_to: str) -> int:
    d_to = (
        datetime.today()
        if date_to == "Present"
        else datetime.strptime(date_to + "-01", "%Y-%m-%d")
    )
    d_fr = datetime.strptime(date_from + "-01", "%Y-%m-%d")
    delta = relativedelta(d_to, d_fr)
    return delta.years * 12 + delta.months


def fmt_cert_expiry(cert: dict) -> str:
    expiry = cert.get("expiry_year")
    if expiry is None:
        return "(perpetual)"
    if expiry < datetime.today().year:
        return f"(expired in {expiry})"
    return f"(expires in {expiry})"


# ─────────────────────────────────────────────
# DOCUMENT SETUP
# ─────────────────────────────────────────────


def setup_document() -> Document:
    doc = Document()
    for section in doc.sections:
        section.top_margin = Cm(2.9)
        section.bottom_margin = Cm(2.6)
        section.left_margin = Cm(1.5)
        section.right_margin = Cm(1.2)
    normal = doc.styles["Normal"]
    normal.font.name = FONT_NAME
    normal.font.size = Pt(DEFAULT_SIZE)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(0)
    return doc


# ─────────────────────────────────────────────
# HEADER & FOOTER
# ─────────────────────────────────────────────


def build_header(doc: Document, first_name: str, last_name: str):
    section = doc.sections[0]
    header = section.header
    for p in header.paragraphs:
        p._element.getparent().remove(p._element)

    tbl = header.add_table(rows=1, cols=3, width=Cm(COL_LEFT_CM + COL_RIGHT_CM))
    set_table_borders_none(tbl)
    row = tbl.rows[0]
    for i, (w, align, text) in enumerate(
        [
            (3.0, WD_ALIGN_PARAGRAPH.LEFT, "Europass"),
            (9.0, WD_ALIGN_PARAGRAPH.CENTER, "Curriculum Vitae"),
            (6.3, WD_ALIGN_PARAGRAPH.RIGHT, f"{first_name} {last_name}"),
        ]
    ):
        set_col_width(row.cells[i], w)
        p = row.cells[i].paragraphs[0]
        p.alignment = align
        set_para_spacing(p, 0, 0)
        add_run(p, text, size=10, color=COLOR_BLUE_LIGHT)


def build_footer(doc: Document):
    section = doc.sections[0]
    footer = section.footer
    for p in footer.paragraphs:
        p._element.getparent().remove(p._element)

    para = footer.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    set_para_spacing(para, 0, 0)

    def _field(name):
        begin = OxmlElement("w:fldChar")
        begin.set(qn("w:fldCharType"), "begin")
        instr = OxmlElement("w:instrText")
        instr.text = name
        end = OxmlElement("w:fldChar")
        end.set(qn("w:fldCharType"), "end")
        for el in [begin, instr, end]:
            r = para.add_run()
            r.font.size = Pt(7)
            r.font.color.rgb = COLOR_BLUE_LIGHT
            r.font.name = FONT_NAME
            r._r.append(el)

    add_run(para, "Page ", size=7, color=COLOR_BLUE_LIGHT)
    _field("PAGE")
    add_run(para, " / ", size=7, color=COLOR_BLUE_LIGHT)
    _field("NUMPAGES")


# ─────────────────────────────────────────────
# TWO-COLUMN HELPERS
# ─────────────────────────────────────────────


def two_col_para(
    doc,
    left_text="",
    right_text="",
    left_size=DEFAULT_SIZE,
    right_size=DEFAULT_SIZE,
    left_color=COLOR_BLUE_DARK,
    right_color=COLOR_BLACK,
    left_bold=False,
    right_bold=False,
    before_pt=0,
    after_pt=0,
):
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    set_para_spacing(para, before_pt, after_pt)
    # Right-aligned tab stop at left column boundary (for label)
    add_tab_stop(para, TABLE_INDENT_CM - 0.3, "right")
    # Left-aligned tab stop for right column content
    add_tab_stop(para, TABLE_INDENT_CM, "left")
    if left_text:
        para.add_run("\t")  # jump to right-aligned tab
        add_run(para, left_text, size=left_size, bold=left_bold, color=left_color)
    para.add_run("\t")  # jump to left-aligned tab for right content
    if right_text:
        add_run(para, right_text, size=right_size, bold=right_bold, color=right_color)
    return para


def right_col_para(
    doc,
    text="",
    size=DEFAULT_SIZE,
    color=COLOR_BLACK,
    bold=False,
    before_pt=0,
    after_pt=0,
    alignment=WD_ALIGN_PARAGRAPH.LEFT,
):
    para = doc.add_paragraph()
    para.alignment = alignment
    set_para_spacing(para, before_pt, after_pt)
    set_para_indent(para, left_cm=TABLE_INDENT_CM)
    if text:
        add_run(para, text, size=size, bold=bold, color=color)
    return para


# ─────────────────────────────────────────────
# BULLET PARAGRAPH (consistent style for project exp & personal skills)
# ─────────────────────────────────────────────


def add_bullet_para(doc, text, before_pt=3, after_pt=0):
    """Add a bullet paragraph aligned to the right column, matching project experience style."""
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    set_para_spacing(para, before_pt, after_pt)
    set_para_indent(
        para,
        left_cm=TABLE_INDENT_CM + BULLET_INDENT_CM,
        hanging_cm=BULLET_INDENT_CM,
    )
    # Bullet char in size 8, text in default size
    add_run(para, BULLET_CHAR, size=8)
    add_run(para, f"  {text}", size=DEFAULT_SIZE)
    return para


# ─────────────────────────────────────────────
# SECTION LABEL + INLINE DIVIDER (consistent alignment)
# ─────────────────────────────────────────────


def section_label_with_divider(doc: Document, label: str):
    """
    Creates a section divider using a two-column table.
    Left cell: label text. Right cell: horizontal line + square.
    Both aligned to the same grid as all other content.
    No cell margins on the divider cells — border sits at text baseline.
    """
    outer = doc.add_table(rows=1, cols=2)
    set_table_borders_none(outer)
    set_table_width(outer, COL_LEFT_CM + COL_RIGHT_CM)
    # No table indent — starts at page left margin like paragraphs
    indent_table(outer, 0)

    row = outer.rows[0]
    c_left = row.cells[0]
    c_right = row.cells[1]

    set_col_width(c_left, COL_LEFT_CM)
    set_col_width(c_right, COL_RIGHT_CM)
    clear_cell_borders(c_left)
    clear_cell_borders(c_right)
    # No margins on outer cells
    set_cell_margins(c_left, 0, 0, 0, 0.3)
    set_cell_margins(c_right, 0, 0, 0, 0)

    # Vertical align bottom so text baseline aligns with divider bottom border
    set_cell_valign(c_left, "bottom")
    set_cell_valign(c_right, "bottom")

    # Label text
    p_lbl = c_left.paragraphs[0]
    p_lbl.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    set_para_spacing(p_lbl, SECTION_SPACE, 0)
    add_run(p_lbl, label.upper(), size=DEFAULT_SIZE, color=COLOR_BLUE_DARK, bold=False)

    # Divider line: single-row inner table with line + square
    inner = c_right.add_table(rows=1, cols=2)
    set_table_borders_none(inner)
    set_table_width(inner, COL_RIGHT_CM)

    i_row = inner.rows[0]
    # Height just enough for the border line — keeps bottom border at text baseline
    set_row_height(i_row, 0.15, exact=True)

    ci_line = i_row.cells[0]
    ci_sq = i_row.cells[1]
    set_col_width(ci_line, COL_RIGHT_CM - 0.25)
    set_col_width(ci_sq, 0.25)
    clear_cell_borders(ci_line)
    clear_cell_borders(ci_sq)

    # NO cell margins on divider cells
    set_cell_margins(ci_line, 0, 0, 0, 0)
    set_cell_margins(ci_sq, 0, 0, 0, 0)

    # Bottom border = the visible line
    set_cell_border(ci_line, bottom={"sz": 8, "val": "single", "color": "398ec9"})
    set_cell_border(ci_sq, bottom={"sz": 8, "val": "single", "color": "398ec9"})
    set_cell_shading(ci_sq, "398ec9")

    # Minimal content so cells render
    for cell in [ci_line, ci_sq]:
        p = cell.paragraphs[0]
        set_para_spacing(p, 0, 0)
        r = p.add_run("")
        r.font.size = Pt(2)


# ─────────────────────────────────────────────
# SECTION: PERSONAL INFORMATION
# ─────────────────────────────────────────────


def build_personal_info(doc: Document, info: dict):
    p_sec = doc.add_paragraph()
    p_sec.alignment = WD_ALIGN_PARAGRAPH.LEFT
    set_para_spacing(p_sec, 0, 4)
    add_tab_stop(p_sec, TABLE_INDENT_CM - 0.3, "right")
    p_sec.add_run("\t")
    add_run(p_sec, "PERSONAL INFORMATION", size=DEFAULT_SIZE, color=COLOR_BLUE_DARK)

    tbl = doc.add_table(rows=1, cols=2)
    set_table_borders_none(tbl)
    row = tbl.rows[0]
    c_pic = row.cells[0]
    c_det = row.cells[1]
    set_col_width(c_pic, COL_LEFT_CM)
    set_col_width(c_det, COL_RIGHT_CM)
    clear_cell_borders(c_pic)
    clear_cell_borders(c_det)
    set_cell_margins(c_pic, 0, 0, 0, 0.19)
    set_cell_margins(c_det, 0, 0, 0, 0)

    p_ph = c_pic.paragraphs[0]
    set_para_spacing(p_ph, 0, 0)
    add_run(p_ph, "[Photo 3.8×3.8cm]", size=DEFAULT_SIZE)

    p_name = c_det.paragraphs[0]
    set_para_spacing(p_name, 0, 3)
    add_run(
        p_name, f"{info['first_name']} {info['last_name']}", size=13, color=COLOR_BLACK
    )

    p_addr = c_det.add_paragraph()
    set_para_spacing(p_addr, 0, 2)
    add_run(p_addr, "\U0001f4cd  ", size=DEFAULT_SIZE)
    add_run(p_addr, info.get("address", ""), size=DEFAULT_SIZE)

    p_phone = c_det.add_paragraph()
    set_para_spacing(p_phone, 0, 2)
    add_run(p_phone, "\U0001f4de  ", size=DEFAULT_SIZE)
    add_run(p_phone, info.get("phone", ""), size=DEFAULT_SIZE)

    p_email = c_det.add_paragraph()
    set_para_spacing(p_email, 0, 2)
    add_run(p_email, "\u2709  ", size=DEFAULT_SIZE)
    email = info.get("email", "")
    if email:
        add_hyperlink(p_email, email, f"mailto:{email}")

    p_sex = c_det.add_paragraph()
    set_para_spacing(p_sex, 4, 0)
    add_run(
        p_sex,
        f"Sex {info.get('sex', '')} | Nationality {info.get('nationality', '')}",
        size=DEFAULT_SIZE,
        color=COLOR_BLUE_LIGHT,
    )


# ─────────────────────────────────────────────
# SECTION: PROPOSED ROLE
# ─────────────────────────────────────────────


def build_proposed_role(doc: Document, tender_info: dict):
    two_col_para(
        doc,
        left_text="PROPOSED ROLE",
        right_text=tender_info.get("proposed_role", ""),
        left_color=COLOR_BLUE_DARK,
        right_size=11,
        right_color=COLOR_BLUE_DARK,
        before_pt=SECTION_SPACE,
    )


# ─────────────────────────────────────────────
# SECTION: PROFESSIONAL EXPERIENCE (tick-box)
# ─────────────────────────────────────────────

BUCKET_LABELS = ["Less than 4", "4 – 9", "10 – 14", "15+"]
BUCKET_KEYS = ["less_than_4", "4_to_9", "10_to_14", "15_plus"]


def build_professional_experience(doc: Document, personal_info: dict):
    """years_experience_bucket is now read from personal_info."""
    two_col_para(
        doc,
        left_text="PROFESSIONAL EXPERIENCE",
        right_text="Number of years of work experience:",
        left_color=COLOR_BLUE_DARK,
        right_size=11,
        right_color=COLOR_BLUE_DARK,
        before_pt=SECTION_SPACE,
        after_pt=4,
    )

    selected = personal_info.get("years_experience_bucket", "")

    tbl = doc.add_table(rows=1, cols=8)
    set_table_borders_none(tbl)
    indent_table(tbl, TABLE_INDENT_CM + 0.19)

    row = tbl.rows[0]
    cells = row.cells
    label_w = 2.5
    tick_w = 0.6

    for i, (lbl, key) in enumerate(zip(BUCKET_LABELS, BUCKET_KEYS)):
        c_lbl = cells[i * 2]
        c_tick = cells[i * 2 + 1]
        set_col_width(c_lbl, label_w)
        set_col_width(c_tick, tick_w)
        set_cell_margins(c_lbl, 0, 0, 0.19, 0.05)
        set_cell_margins(c_tick, 0, 0, 0.05, 0.19)

        p_lbl = c_lbl.paragraphs[0]
        set_para_spacing(p_lbl, 0, 0)
        p_lbl.alignment = WD_ALIGN_PARAGRAPH.LEFT
        add_run(p_lbl, lbl, size=11, color=COLOR_BLUE_DARK)

        p_tick = c_tick.paragraphs[0]
        set_para_spacing(p_tick, 0, 0)
        p_tick.alignment = WD_ALIGN_PARAGRAPH.CENTER
        tick_char = "X" if key == selected else ""
        add_run(p_tick, tick_char, size=11, color=COLOR_BLUE_DARK)

    apply_standard_table_borders(tbl)


# ─────────────────────────────────────────────
# SECTION: WORK EXPERIENCE
# ─────────────────────────────────────────────


def build_work_experience(doc: Document, work_exp: list):
    two_col_para(
        doc,
        left_text="WORK EXPERIENCE",
        left_color=COLOR_BLUE_DARK,
        before_pt=SECTION_SPACE,
        after_pt=2,
    )

    col_widths = [3.5, 5.5, 4.3]
    headers = ["DATE", "TITLE", "ORGANISATION"]

    tbl = doc.add_table(rows=1, cols=3)
    set_table_borders_none(tbl)
    indent_table(tbl, TABLE_INDENT_CM + 0.19)

    hdr_row = tbl.rows[0]
    for i, (hdr, w) in enumerate(zip(headers, col_widths)):
        cell = hdr_row.cells[i]
        set_col_width(cell, w)
        set_cell_margins(cell, 0, 0, 0.19, 0.19)
        p = cell.paragraphs[0]
        set_para_spacing(p, 0, 0)
        add_run(p, hdr, size=DEFAULT_SIZE, bold=True, color=COLOR_BLUE_DARK)

    for entry in work_exp:
        row = tbl.add_row()
        date_str = fmt_date_range(entry["date_from"], entry["date_to"])
        org_str = entry["organisation"]
        if entry.get("organisation_country"):
            org_str += f", {entry['organisation_country']}"
        values = [date_str, entry["job_title"], org_str]
        for i, (val, w) in enumerate(zip(values, col_widths)):
            cell = row.cells[i]
            set_col_width(cell, w)
            set_cell_margins(cell, 0, 0, 0.19, 0.19)
            p = cell.paragraphs[0]
            set_para_spacing(p, 0, 0)
            add_run(p, val, size=DEFAULT_SIZE, color=COLOR_BLACK)

    apply_standard_table_borders(tbl)


# ─────────────────────────────────────────────
# SECTION: PROFILE
# ─────────────────────────────────────────────


def build_profile(doc: Document, profile_text: str):
    two_col_para(
        doc,
        left_text="PROFILE",
        left_color=COLOR_BLUE_DARK,
        before_pt=SECTION_SPACE,
        after_pt=2,
    )

    for para_text in profile_text.strip().split("\n"):
        para_text = para_text.strip()
        if not para_text:
            continue
        right_col_para(
            doc, text=para_text, after_pt=4, alignment=WD_ALIGN_PARAGRAPH.JUSTIFY
        )


# ─────────────────────────────────────────────
# FULL-WIDTH TABLES
# ─────────────────────────────────────────────


def _full_width_label(doc: Document, label: str):
    para = doc.add_paragraph()
    set_para_spacing(para, SECTION_SPACE, 6)  # 6pt gap between label and table
    add_run(
        para, label, size=DEFAULT_SIZE, bold=True, underline=True, color=COLOR_BLACK
    )


def _style_full_table_header(row, col_widths: list):
    for i, cell in enumerate(row.cells):
        set_col_width(cell, col_widths[i])
        set_cell_margins(cell, 0, 0, 0.19, 0.19)
        set_cell_shading(cell, "156082")
        p = cell.paragraphs[0]
        set_para_spacing(p, 0, 0)
        for run in p.runs:
            apply_run_fmt(
                run, size=DEFAULT_SIZE, bold=True, color=RGBColor(0xFF, 0xFF, 0xFF)
            )


def build_requirements_matrix(doc: Document, tender_info: dict):
    reqs = tender_info.get("requirements_matrix", [])
    full_name = (
        tender_info.get("_first_name", "") + " " + tender_info.get("_last_name", "")
    ).strip()

    _full_width_label(doc, "Requirements matrix")

    col_widths = [14.5, 3.8]
    tbl = doc.add_table(rows=1, cols=2)
    set_table_borders_none(tbl)
    set_table_width(tbl, FULL_WIDTH_CM)

    hdr = tbl.rows[0]
    hdr.cells[0].paragraphs[0].add_run("Requirement")
    hdr.cells[1].paragraphs[0].add_run(full_name)
    _style_full_table_header(hdr, col_widths)

    for req in reqs:
        row = tbl.add_row()
        for i, w in enumerate(col_widths):
            set_col_width(row.cells[i], w)
            set_cell_margins(row.cells[i], 0, 0, 0.19, 0.19)

        p0 = row.cells[0].paragraphs[0]
        set_para_spacing(p0, 0, 0)
        p0.alignment = WD_ALIGN_PARAGRAPH.LEFT
        add_run(p0, req["requirement_text"], size=DEFAULT_SIZE)

        p1 = row.cells[1].paragraphs[0]
        set_para_spacing(p1, 0, 0)
        p1.alignment = WD_ALIGN_PARAGRAPH.LEFT
        add_run(p1, "☒", size=DEFAULT_SIZE)

    apply_standard_table_borders(tbl)


def build_experience_overview(doc: Document, overview: list):
    _full_width_label(doc, "Experience overview")

    col_widths = [2.5, 11.5, 1.8, 2.5]
    headers = [
        "Date",
        "Roles and Responsibilities",
        "Number of Months",
        "Total Number of Months of Relevant Professional Experience",
    ]

    tbl = doc.add_table(rows=1, cols=4)
    set_table_borders_none(tbl)
    set_table_width(tbl, FULL_WIDTH_CM)

    # Force layout type to fixed so Word respects our widths
    tblPr = tbl._tbl.find(qn("w:tblPr"))
    if tblPr is None:
        tblPr = OxmlElement("w:tblPr")
        tbl._tbl.insert(0, tblPr)
    for existing in tblPr.findall(qn("w:tblLayout")):
        tblPr.remove(existing)
    tblLayout = OxmlElement("w:tblLayout")
    tblLayout.set(qn("w:type"), "fixed")
    tblPr.append(tblLayout)

    hdr = tbl.rows[0]
    for i, h in enumerate(headers):
        hdr.cells[i].paragraphs[0].add_run(h)
    _style_full_table_header(hdr, col_widths)

    for entry in overview:
        row = tbl.add_row()
        duration = entry.get("duration_months") or calc_months(
            entry["date_from"], entry["date_to"]
        )
        relevant = entry.get("relevant_months", duration)

        date_str = fmt_date_range(entry["date_from"], entry["date_to"])

        for i, w in enumerate(col_widths):
            set_col_width(row.cells[i], w)
            set_cell_margins(row.cells[i], 0, 0, 0.19, 0.19)

        p0 = row.cells[0].paragraphs[0]
        set_para_spacing(p0, 0, 0)
        add_run(p0, date_str, size=DEFAULT_SIZE)

        p1 = row.cells[1].paragraphs[0]
        set_para_spacing(p1, 0, 0)
        add_run(p1, "Role: ", size=DEFAULT_SIZE, bold=True)
        add_run(p1, entry["role"], size=DEFAULT_SIZE)
        add_run(p1, "\nResponsibilities: ", size=DEFAULT_SIZE, bold=True)
        add_run(p1, entry["responsibilities"], size=DEFAULT_SIZE)

        p2 = row.cells[2].paragraphs[0]
        set_para_spacing(p2, 0, 0)
        add_run(p2, str(duration), size=DEFAULT_SIZE)

        p3 = row.cells[3].paragraphs[0]
        set_para_spacing(p3, 0, 0)
        add_run(p3, str(relevant), size=DEFAULT_SIZE)

    apply_standard_table_borders(tbl)

    spacer = doc.add_paragraph()
    set_para_spacing(spacer, 0, 8)


# ─────────────────────────────────────────────
# SECTION: PROJECT EXPERIENCE
# ─────────────────────────────────────────────


def build_project_experience(doc: Document, projects: list):
    section_label_with_divider(doc, "Project Experience")

    for proj in projects:
        date_str = fmt_date_range(proj["date_from"], proj["date_to"])

        two_col_para(
            doc,
            left_text=date_str,
            right_text=proj.get("role", ""),
            left_color=COLOR_BLUE_DARK,
            right_size=11,
            right_color=COLOR_BLUE_DARK,
            before_pt=6,
        )

        sector_parts = [p for p in [proj.get("sector"), proj.get("country")] if p]
        if sector_parts:
            right_col_para(
                doc,
                text=", ".join(sector_parts),
                color=COLOR_BLUE_LIGHT,
                before_pt=4,
                after_pt=2,
            )

        # Bullets — using shared helper
        for bullet in proj.get("bullets", []):
            add_bullet_para(doc, bullet, before_pt=3, after_pt=0)

        # Allocation
        alloc = proj.get("allocation_percent")
        if alloc:
            p_alloc = right_col_para(doc, before_pt=3, after_pt=0)
            add_run(p_alloc, "Project Allocation in % (only specified in case of parallel projects)", size=DEFAULT_SIZE, underline=True)
            add_run(p_alloc, f": {alloc}", size=DEFAULT_SIZE)

        spacer = doc.add_paragraph()
        set_para_spacing(spacer, 0, 4)


# ─────────────────────────────────────────────
# SECTION: EDUCATION AND TRAINING
# ─────────────────────────────────────────────


def build_education(doc: Document, education: list):
    section_label_with_divider(doc, "Education and Training")

    for edu in education:
        date_str = fmt_date_range(edu["date_from"], edu["date_to"])

        two_col_para(
            doc,
            left_text=date_str,
            right_text=edu.get("qualification_title", ""),
            left_color=COLOR_BLUE_DARK,
            right_size=11,
            right_color=COLOR_BLUE_DARK,
            before_pt=6,
        )

        inst_parts = [
            p for p in [edu.get("institution"), edu.get("institution_country")] if p
        ]
        if inst_parts:
            right_col_para(
                doc,
                text=", ".join(inst_parts),
                color=COLOR_BLACK,
                before_pt=4,
                after_pt=4,
            )


# ─────────────────────────────────────────────
# SECTION: PERSONAL SKILLS — LANGUAGE TABLE
# ─────────────────────────────────────────────


def build_language_table(doc: Document, languages: dict):
    other = languages.get("other", [])
    if not other:
        return

    two_col_para(
        doc,
        left_text="Other language(s)",
        left_color=COLOR_BLUE_DARK,
        before_pt=4,
        after_pt=2,
    )

    data_col_w = round(COL_RIGHT_CM / 5, 2)
    name_col_w = COL_LEFT_CM + 0.19

    n_rows = 2 + len(other)
    tbl = doc.add_table(rows=n_rows, cols=6)
    set_table_borders_none(tbl)

    GRAY = "BFBFBF"
    BORDER_1PT = {"sz": 8, "val": "single", "color": GRAY}
    BORDER_HALF = {"sz": 4, "val": "single", "color": GRAY}
    NO_BORDER = {"sz": 0, "val": "none", "color": "FFFFFF"}

    def _set_lang_cell(
        cell, w, top=None, bottom=None, left=None, right=None, no_border=False
    ):
        set_col_width(cell, w)
        set_cell_margins(cell, 0, 0, 0.19, 0.19)
        if no_border:
            clear_cell_borders(cell)
        else:
            kwargs = {}
            if top:
                kwargs["top"] = top
            if bottom:
                kwargs["bottom"] = bottom
            if left:
                kwargs["left"] = left
            if right:
                kwargs["right"] = right
            if kwargs:
                set_cell_border(cell, **kwargs)
            else:
                clear_cell_borders(cell)

    # Row 0: group headers
    r0 = tbl.rows[0]
    _set_lang_cell(r0.cells[0], name_col_w, no_border=True)
    p = r0.cells[0].paragraphs[0]
    set_para_spacing(p, 0, 0)

    r0.cells[1].merge(r0.cells[2])
    r0.cells[3].merge(r0.cells[4])

    group_labels = {1: "UNDERSTANDING", 3: "SPEAKING", 5: "WRITING"}
    for col_idx in [1, 3, 5]:
        cell = r0.cells[col_idx]
        set_col_width(cell, data_col_w * (2 if col_idx in [1, 3] else 1))
        set_cell_margins(cell, 0, 0, 0.19, 0.19)
        set_cell_border(
            cell,
            bottom=BORDER_1PT,
            left=BORDER_1PT if col_idx in [3, 5] else NO_BORDER,
        )
        p = cell.paragraphs[0]
        set_para_spacing(p, 0, 0)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        add_run(p, group_labels[col_idx], size=7, color=COLOR_BLUE_DARK)

    # Row 1: sub-headers
    r1 = tbl.rows[1]
    sub_labels = [
        "",
        "Listening",
        "Reading",
        "Spoken interaction",
        "Spoken production",
        "Writing",
    ]
    for i, lbl in enumerate(sub_labels):
        cell = r1.cells[i]
        w = name_col_w if i == 0 else data_col_w
        set_col_width(cell, w)
        set_cell_margins(cell, 0, 0, 0.19, 0.19)
        if i == 0:
            clear_cell_borders(cell)
        else:
            set_cell_border(
                cell,
                bottom=BORDER_1PT,
                left=(
                    BORDER_1PT if i in [3, 5] else BORDER_HALF if i == 2 else NO_BORDER
                ),
            )
        p = cell.paragraphs[0]
        set_para_spacing(p, 0, 0)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        if lbl:
            add_run(p, lbl, size=8, color=COLOR_BLUE_DARK)

    # Data rows
    for j, lang in enumerate(other):
        row = tbl.rows[2 + j]
        is_last = j == len(other) - 1
        vals = [
            lang["language"],
            lang["listening"],
            lang["reading"],
            lang["spoken_interaction"],
            lang["spoken_production"],
            lang["writing"],
        ]
        for i, val in enumerate(vals):
            cell = row.cells[i]
            w = name_col_w if i == 0 else data_col_w
            set_col_width(cell, w)
            set_cell_margins(cell, 0, 0, 0.19, 0.19)

            if i == 0:
                clear_cell_borders(cell)
                p = cell.paragraphs[0]
                p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            else:
                border_kwargs = {}
                if not is_last:
                    border_kwargs["bottom"] = BORDER_HALF
                if i in [3, 5]:
                    border_kwargs["left"] = BORDER_1PT
                elif i == 2:
                    border_kwargs["left"] = BORDER_HALF
                if border_kwargs:
                    set_cell_border(cell, **border_kwargs)
                else:
                    clear_cell_borders(cell)
                p = cell.paragraphs[0]
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER

            set_para_spacing(p, 0, 0)
            add_run(p, val, size=DEFAULT_SIZE, color=COLOR_BLACK)

    # CEFR disclaimer
    for line in [
        "Levels: A1/2: Basic user - B1/2: Independent user - C1/2 Proficient user",
        "Common European Framework of Reference for Languages",
    ]:
        p = doc.add_paragraph()
        set_para_spacing(p, 0, 0)
        set_para_indent(p, left_cm=TABLE_INDENT_CM)
        add_run(p, line, size=7.5, color=COLOR_BLUE_DARK)


# ─────────────────────────────────────────────
# SECTION: PERSONAL SKILLS
# ─────────────────────────────────────────────


def _skill_bullets(doc, label, bullets, before_pt=SKILL_SUBSECTION_SPACE):
    """Render a personal-skills subsection with bullets matching project experience style."""
    if not bullets:
        return

    # First bullet on same line as label (via tab)
    para0 = doc.add_paragraph()
    para0.alignment = WD_ALIGN_PARAGRAPH.LEFT
    set_para_spacing(para0, before_pt, 0)
    add_tab_stop(para0, TABLE_INDENT_CM, "left")
    set_para_indent(
        para0,
        left_cm=TABLE_INDENT_CM + BULLET_INDENT_CM,
        hanging_cm=TABLE_INDENT_CM + BULLET_INDENT_CM,
    )
    add_run(para0, label, size=DEFAULT_SIZE, color=COLOR_BLUE_DARK)
    para0.add_run("\t")
    add_run(para0, BULLET_CHAR, size=8)
    add_run(para0, f"  {bullets[0]}", size=DEFAULT_SIZE)

    # Remaining bullets — same style as project experience bullets
    for bullet in bullets[1:]:
        add_bullet_para(doc, bullet, before_pt=3, after_pt=0)


def build_personal_skills(doc: Document, skills: dict, languages: dict):
    section_label_with_divider(doc, "Personal Skills")

    # Mother tongue
    mother = languages.get("mother_tongue", [])
    two_col_para(
        doc,
        left_text="Mother tongue(s)",
        right_text=", ".join(mother),
        left_color=COLOR_BLUE_DARK,
        before_pt=4,
    )

    # Language table
    build_language_table(doc, languages)

    # Skill subsections with consistent spacing between them
    _skill_bullets(
        doc,
        "Communication skills",
        skills.get("communication", []),
        before_pt=SKILL_SUBSECTION_SPACE,
    )

    _skill_bullets(
        doc,
        "Organisational / managerial skills",
        skills.get("organisational_managerial", []),
        before_pt=SKILL_SUBSECTION_SPACE,
    )

    _skill_bullets(
        doc,
        "Computer skills",
        skills.get("computer_skills", []),
        before_pt=SKILL_SUBSECTION_SPACE,
    )

    # Certifications
    certs = skills.get("certifications", [])
    if certs:
        # First cert with label
        cert = certs[0]
        expiry_str = fmt_cert_expiry(cert)
        para = doc.add_paragraph()
        para.alignment = WD_ALIGN_PARAGRAPH.LEFT
        set_para_spacing(para, SKILL_SUBSECTION_SPACE, 0)
        add_tab_stop(para, TABLE_INDENT_CM, "left")
        set_para_indent(
            para,
            left_cm=TABLE_INDENT_CM + BULLET_INDENT_CM,
            hanging_cm=TABLE_INDENT_CM + BULLET_INDENT_CM,
        )
        add_run(para, "Certifications", size=DEFAULT_SIZE, color=COLOR_BLUE_DARK)
        para.add_run("\t")
        add_run(para, BULLET_CHAR, size=8)
        add_run(para, f"  {cert['year']} \u2013 {cert['title']} ", size=DEFAULT_SIZE)
        add_run(para, expiry_str, size=7, color=COLOR_BLACK)

        # Remaining certs
        for cert in certs[1:]:
            expiry_str = fmt_cert_expiry(cert)
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            set_para_spacing(p, 3, 0)
            set_para_indent(
                p,
                left_cm=TABLE_INDENT_CM + BULLET_INDENT_CM,
                hanging_cm=BULLET_INDENT_CM,
            )
            add_run(p, BULLET_CHAR, size=8)
            add_run(p, f"  {cert['year']} \u2013 {cert['title']} ", size=DEFAULT_SIZE)
            add_run(p, expiry_str, size=7, color=COLOR_BLACK)


# ─────────────────────────────────────────────
# MAIN GENERATOR
# ─────────────────────────────────────────────


def generate_cv(cv_json: dict, output_path: str):
    doc = setup_document()

    pi = cv_json["personal_info"]
    tender_info = cv_json.get("tender_info", {})
    tender_info["_first_name"] = pi["first_name"]
    tender_info["_last_name"] = pi["last_name"]

    build_header(doc, pi["first_name"], pi["last_name"])
    build_footer(doc)

    build_personal_info(doc, pi)
    build_proposed_role(doc, tender_info)
    build_professional_experience(doc, pi)  # ← Now reads from personal_info
    build_work_experience(doc, cv_json.get("work_experience", []))
    build_profile(doc, cv_json.get("profile", ""))
    build_requirements_matrix(doc, tender_info)
    build_experience_overview(doc, cv_json.get("experience_overview", []))
    build_project_experience(doc, cv_json.get("project_experience", []))
    build_education(doc, cv_json.get("education", []))
    build_personal_skills(
        doc, cv_json.get("personal_skills", {}), cv_json.get("languages", {})
    )

    doc.save(output_path)
    print(f"✓ CV saved → {output_path}")


# ─────────────────────────────────────────────
# ENTRY POINT
# ─────────────────────────────────────────────

if __name__ == "__main__":
    if len(sys.argv) < 3:
        print("Usage: python cv_generator.py input.json output.docx")
        sys.exit(1)
    with open(sys.argv[1], "r", encoding="utf-8") as f:
        data = json.load(f)
    generate_cv(data, sys.argv[2])
    import os

    os.startfile(sys.argv[2])
